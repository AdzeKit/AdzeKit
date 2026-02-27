"""Markdown-to-docx export via Pandoc.

Converts any markdown file in the shed to .docx format using a
custom reference document styled with DM Sans. Requires pandoc to
be installed and available on PATH.
"""

import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Emu
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ---------------------------------------------------------------------------
# Locate bundled assets shipped inside the package
# ---------------------------------------------------------------------------

def _assets_dir() -> Path:
    """Return the path to the package's assets directory."""
    return Path(__file__).resolve().parents[3] / "assets"


# Font configuration
_FONT_NAME = "DM Sans"
_FONT_COLOR = RGBColor(0x1A, 0x1A, 0x1A)


# ---------------------------------------------------------------------------
# Reference document (controls all pandoc docx styles)
# ---------------------------------------------------------------------------

def _build_reference_doc(ref_path: Path) -> Path:
    """Create a minimal reference.docx with DM Sans styles.

    Pandoc uses a reference document to derive all styles for the
    generated docx.  We build one programmatically so there is no
    opaque binary blob checked into the repo.
    """
    doc = Document()

    # Remove all default section breaks -- single continuous section
    for section in doc.sections:
        section.start_type = 0  # CONTINUOUS

    # Apply DM Sans to every built-in style
    for style in doc.styles:
        if hasattr(style, "font") and style.font is not None:
            style.font.name = _FONT_NAME
            style.font.color.rgb = _FONT_COLOR

        if hasattr(style, "paragraph_format"):
            pf = style.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(4)

    # Heading styles -- clean, tight spacing
    sizes = {1: 18, 2: 14, 3: 12, 4: 11}
    for level, size in sizes.items():
        name = f"Heading {level}"
        if name in doc.styles:
            s = doc.styles[name]
            s.font.bold = True
            s.font.size = Pt(size)
            s.font.color.rgb = _FONT_COLOR
            s.paragraph_format.space_before = Pt(10 if level == 1 else 8)
            s.paragraph_format.space_after = Pt(2)
            s.paragraph_format.line_spacing = 1.0

    # Normal body text
    normal = doc.styles["Normal"]
    normal.font.size = Pt(10)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(3)

    doc.save(str(ref_path))
    return ref_path


def _get_reference_doc() -> Path:
    """Return path to the cached reference.docx, building it if needed."""
    ref = _assets_dir() / "reference.docx"
    # Always rebuild so style changes take effect
    _build_reference_doc(ref)
    return ref


# ---------------------------------------------------------------------------
# Post-processing: table borders and section break removal
# ---------------------------------------------------------------------------

_BORDER_XML = (
    '<w:tcBorders %s>'
    '  <w:top w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
    '  <w:left w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
    '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
    '  <w:right w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
    '</w:tcBorders>'
) % nsdecls("w")


def _style_tables(doc: Document) -> None:
    """Add clean borders to every cell in every table."""
    for table in doc.tables:
        # Remove any table-level style that pandoc may have applied
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        if tbl_pr is None:
            tbl_pr = parse_xml(f"<w:tblPr {nsdecls('w')}/>")
            tbl.insert(0, tbl_pr)

        # Set table-level borders
        borders_el = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            '  <w:top w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
            '  <w:left w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
            '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
            '  <w:right w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
            '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
            '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
            "</w:tblBorders>"
        )
        # Remove existing borders if any
        existing = tbl_pr.find(qn("w:tblBorders"))
        if existing is not None:
            tbl_pr.remove(existing)
        tbl_pr.append(borders_el)

        # Cell-level borders for completeness
        for row in table.rows:
            for cell in row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                existing_borders = tc_pr.find(qn("w:tcBorders"))
                if existing_borders is not None:
                    tc_pr.remove(existing_borders)
                tc_pr.append(parse_xml(_BORDER_XML))

                # Set cell font
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = _FONT_NAME
                        run.font.size = Pt(9)


def _remove_section_breaks(doc: Document) -> None:
    """Convert all section breaks to continuous so the doc flows."""
    for section in doc.sections:
        sectPr = section._sectPr
        sect_type = sectPr.find(qn("w:type"))
        if sect_type is not None:
            sect_type.set(qn("w:val"), "continuous")
        else:
            sectPr.append(
                parse_xml(f'<w:type {nsdecls("w")} w:val="continuous"/>')
            )


def _postprocess(docx_path: Path) -> None:
    """Apply table borders and remove section breaks from a docx."""
    doc = Document(str(docx_path))
    _style_tables(doc)
    _remove_section_breaks(doc)
    doc.save(str(docx_path))


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def _check_pandoc() -> None:
    """Raise if pandoc is not installed."""
    if shutil.which("pandoc") is None:
        raise RuntimeError(
            "pandoc is not installed. Install it with: brew install pandoc"
        )


def to_docx(source: Path, output: Path | None = None) -> Path:
    """Convert a markdown file to .docx using pandoc.

    Uses a DM Sans reference document for styling, adds clean table
    borders, and removes section breaks.  If output is not specified,
    the docx is written alongside the source with the same stem.
    """
    _check_pandoc()

    if not source.exists():
        raise FileNotFoundError(f"Source file not found: {source}")

    if output is None:
        output = source.with_suffix(".docx")

    ref_doc = _get_reference_doc()

    subprocess.run(
        [
            "pandoc",
            str(source),
            "-f", "markdown",
            "-t", "docx",
            "-o", str(output),
            "--standalone",
            f"--reference-doc={ref_doc}",
        ],
        check=True,
        capture_output=True,
        text=True,
    )

    _postprocess(output)

    return output
